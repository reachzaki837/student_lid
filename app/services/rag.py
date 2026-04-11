import os
import io
import base64
import networkx as nx
from typing import List, Dict, Any, Optional
from PyPDF2 import PdfReader
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from langchain_core.messages import HumanMessage
from app.models.user import Material
from app.core.config import settings

# Vercel serverless functions strictly only allow writing to the /tmp directory.
INDEX_DIR = "/tmp/faiss_index"
os.makedirs(INDEX_DIR, exist_ok=True)

# Google Gemini Embeddings (API-based, lightweight deployment)
# Important: This avoids the massive ~5GB sentence-transformers dependency for Vercel.
# Try current embedding model first, then fall back for older projects/keys.
EMBEDDING_MODELS = tuple(dict.fromkeys([
    os.getenv("GOOGLE_EMBEDDING_MODEL", "models/text-embedding-004"),
    "models/embedding-001",
]))
_ACTIVE_EMBEDDING_MODEL: Optional[str] = None
_ACTIVE_EMBEDDINGS: Optional[GoogleGenerativeAIEmbeddings] = None


def _is_embedding_not_found_error(exc: Exception) -> bool:
    message = str(exc).lower()
    return "404" in message and ("not_found" in message or "not found" in message)


def _set_active_embeddings(model_name: str) -> GoogleGenerativeAIEmbeddings:
    global _ACTIVE_EMBEDDING_MODEL, _ACTIVE_EMBEDDINGS
    _ACTIVE_EMBEDDING_MODEL = model_name
    _ACTIVE_EMBEDDINGS = GoogleGenerativeAIEmbeddings(
        model=model_name,
        google_api_key=settings.GOOGLE_API_KEY,
    )
    return _ACTIVE_EMBEDDINGS


def _get_active_embeddings() -> GoogleGenerativeAIEmbeddings:
    if _ACTIVE_EMBEDDINGS is None:
        return _set_active_embeddings(EMBEDDING_MODELS[0])
    return _ACTIVE_EMBEDDINGS


def _switch_to_fallback_embedding_model() -> bool:
    current = _ACTIVE_EMBEDDING_MODEL or EMBEDDING_MODELS[0]
    for candidate in EMBEDDING_MODELS:
        if candidate != current:
            _set_active_embeddings(candidate)
            return True
    return False

SUPPORTED_UPLOAD_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".bmp",
    ".gif",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}

IMAGE_MIME_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".gif": "image/gif",
}

class RAGService:
    @staticmethod
    async def _extract_text_from_image(file_bytes: bytes, filename: str) -> str:
        """Use Gemini vision to extract educational text and context from uploaded images."""
        if not settings.GOOGLE_API_KEY:
            raise ValueError("GOOGLE_API_KEY is required for image analysis uploads")

        extension = os.path.splitext(filename.lower())[1]
        mime_type = IMAGE_MIME_TYPES.get(extension, "image/jpeg")
        encoded_image = base64.b64encode(file_bytes).decode("ascii")

        from langchain_google_genai import ChatGoogleGenerativeAI

        llm = ChatGoogleGenerativeAI(
            temperature=0,
            model="gemini-2.5-flash",
            google_api_key=settings.GOOGLE_API_KEY,
        )

        response = await llm.ainvoke([
            HumanMessage(content=[
                {
                    "type": "text",
                    "text": (
                        "Analyze this uploaded classroom image and extract useful text/context for retrieval. "
                        "Return concise, plain text containing: 1) OCR text, 2) key topic summary, "
                        "3) important formulas/definitions, 4) likely question-answer cues for teaching."
                    ),
                },
                {
                    "type": "image_url",
                    "image_url": f"data:{mime_type};base64,{encoded_image}",
                },
            ])
        ])

        content = response.content
        if isinstance(content, str):
            return content.strip()

        if isinstance(content, list):
            text_parts = []
            for part in content:
                if isinstance(part, str):
                    text_parts.append(part)
                elif isinstance(part, dict) and part.get("text"):
                    text_parts.append(str(part["text"]))
            return "\n".join(text_parts).strip()

        return str(content).strip()

    @staticmethod
    async def process_upload(file_bytes: bytes, filename: str, uploader_email: str) -> tuple[bool, str]:
        """Processes an uploaded file, extracts text, chunks it, and adds to FAISS index."""
        try:
            if not settings.GOOGLE_API_KEY:
                return False, "GOOGLE_API_KEY is required for document uploads."

            extension = os.path.splitext(filename.lower())[1]
            if extension not in SUPPORTED_UPLOAD_EXTENSIONS:
                supported = ", ".join(sorted(SUPPORTED_UPLOAD_EXTENSIONS))
                return False, f"Unsupported file type '{extension}'. Supported types: {supported}."

            # 1. Extract Text
            content = ""
            if extension == ".pdf":
                pdf = PdfReader(io.BytesIO(file_bytes))
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
            elif extension == ".docx":
                try:
                    from docx import Document as DocxDocument
                except ImportError as exc:
                    raise ValueError("python-docx is required to process .docx files") from exc

                doc = DocxDocument(io.BytesIO(file_bytes))
                content = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
            elif extension in IMAGE_EXTENSIONS:
                content = await RAGService._extract_text_from_image(file_bytes, filename)
            else:
                content = file_bytes.decode("utf-8", errors="ignore")

            if not content or not content.strip():
                return False, "No readable content found in this file. Please upload a clearer file."
            
            # Save the material to MongoDB
            material = Material(uploader_email=uploader_email, filename=filename, content_text=content)
            await material.insert()

            # 2. Chunk the text
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            chunks = text_splitter.split_text(content)
            
            docs = [
                LangchainDocument(
                    page_content=chunk,
                    metadata={"source": filename, "uploader": uploader_email}
                ) for chunk in chunks
            ]
            
            # 3. Add to FAISS Vector Store
            def _upsert_docs_with_embeddings(emb: GoogleGenerativeAIEmbeddings) -> None:
                try:
                    vectorstore = FAISS.load_local(INDEX_DIR, emb, allow_dangerous_deserialization=True)
                    vectorstore.add_documents(docs)
                    vectorstore.save_local(INDEX_DIR)
                except Exception:
                    # Create a new index if it doesn't exist
                    vectorstore = FAISS.from_documents(docs, emb)
                    vectorstore.save_local(INDEX_DIR)

            try:
                _upsert_docs_with_embeddings(_get_active_embeddings())
            except Exception as exc:
                if _is_embedding_not_found_error(exc) and _switch_to_fallback_embedding_model():
                    _upsert_docs_with_embeddings(_get_active_embeddings())
                else:
                    raise
                
            return True, "ok"
        except Exception as e:
            print(f"Error processing document: {e}")
            return False, str(e)

    @staticmethod
    def query_documents(query: str, top_k: int = 3) -> str:
        """Queries the FAISS index to act as the retriever for GraphRAG/Standard RAG."""
        try:
            if not os.path.exists(os.path.join(INDEX_DIR, "index.faiss")):
                return "No specific course materials found."

            def _query_with_embeddings(emb: GoogleGenerativeAIEmbeddings) -> str:
                vectorstore = FAISS.load_local(INDEX_DIR, emb, allow_dangerous_deserialization=True)
                results = vectorstore.similarity_search(query, k=top_k)
                return "\n\n".join(
                    [f"Source: {res.metadata.get('source', 'Unknown')}\nContent: {res.page_content}" for res in results]
                )

            try:
                context = _query_with_embeddings(_get_active_embeddings())
            except Exception as exc:
                if _is_embedding_not_found_error(exc) and _switch_to_fallback_embedding_model():
                    context = _query_with_embeddings(_get_active_embeddings())
                else:
                    raise

            return context
        except Exception as e:
            print(f"FAISS Query Error: {e}")
            return "No specific course materials found."
